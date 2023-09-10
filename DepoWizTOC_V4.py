import streamlit as st
import requests
import PyPDF2
import json
import pandas as pd
from io import BytesIO
from docx import Document
import base64
import re
from docx.shared import Inches
import io
from collections import Counter
import numpy as np
from DepoWizQNA_V3 import make_valid_partition_name, save_to_database, extract_pdf_text_to_dataframe
import nltk
from nltk.corpus import stopwords
from builtins import PendingDeprecationWarning



# My API Key
api_key = "YOUR_API_KEY"

token = "50.0"
temp = "0.0"
top_p = "0.90"
f_pen = "0.0"
p_pen = "0.0"
 
nltk.download('stopwords')
nltk.download('punkt')

def get_completions(prompt, temp='0.0',top_p='0.75'):
    url = "https://chat-gpt-a1.openai.azure.com/openai/deployments/DanielChatGPT/chat/completions?api-version=2023-03-15-preview"
    headers = {'Content-Type': 'application/json',
               'api-key': api_key,
               'max_tokens': str(token),
               'temperature': str(temp) if temp else "0.0",
               'top_p': str(top_p) if top_p else "0.75",
               'frequency_penalty': str(f_pen) if f_pen else "0.0",
               'Presence_penalty': str(p_pen) if p_pen else "0.0"
               }
    data = prompt

    body = str.encode(json.dumps(data))

    r = requests.post(url=url, data=body, headers=headers)
    response = r.content

    response = json.loads(response.decode())
    if not ('choices' in response):
        return
    # print(response)
    response = response["choices"][0]["message"]["content"]
    
    return response

# ****************************** Functions for Generating TOC *************************************************


def toc_generator(text,gender):
    prompt = {'messages': [{"role": "system", "content": f"""You are a Table of Content generator by extracting context from deposition document and converting it to telegram text. Remember the gender of witness:
    Witness's Gender: {gender}"""},

                         {"role": "user", "content": f"""
You are a Table of Content generator for the deposition document.
You are provided with the text extracted from deposition pdf delimited by ``` .
Your task is to perform the following actions for the topics disscused in the deposition:
1. Generate a list of overview Titles for the discussion.
2. For each title extract the page number.
3. For each title extract the list of topics (all possible topics) separated by comma covered in the given deposition.
4. For each title write a brief point to point summary of the conversation.
5. Format it in a list of json objects with following keys (in given Order):
    
    "Page_no": <Page number (Integer datatype) >,
    "Title": <Generated Title (Overview of the Discussion)>,
    "Topics": <List of topics (comma separated) discussed in this title.>,
    "Summary": <Brief point to point summary of the conversation(bullet points) >

6. Give Output in a list of json objects.

Given Text: ```{text}```

"""}]}
    
    return get_completions(prompt)


def fromatString(text):
    prompt = {'messages': [{"role": "system", "content": "You are a helper function to remove the syntax error from the given string and returns the formatted output."},
                         {"role": "user", "content": f"""
You are provided the info delimited by ```.
Your task is to perform the following actions:
1. Remove the syntax error from the given info.
2. Format it properly and return it as an Output.

Given Info: ```{text}```
"""}]}
    result = text
    try:
        result = get_completions(prompt)
    except:
        return result
    return result


def pdfParser(pdf_file,gender):
    if pdf_file:
        indexDict = []
        uploaded_file_name = f"{pdf_file.name}_{pdf_file.size}.pdf"
        with open(uploaded_file_name, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            pdf_text = ''
            count = 0
            # for page_num in range(0,len(pdf_reader.pages)):
            pgSize = len(pdf_reader.pages)
            value_slot = st.empty()

            for page_num in range(0,pgSize):
            
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if not text:
                    # st.warning("Unable to Parse Pdf. Skipping Page")
                    continue
                else:
                    value_slot.markdown(f"Processing Pages:  {page_num+1}/{pgSize}")
                    pdf_text = pdf_text + f'***** Page no: {page_num+1} ***** ' + text + f"\n\t\t ********* page end ********** \n"
                    count += 1
                    if count == 3:
                        try:
                            index = extract_json_objects(toc_generator(pdf_text,gender))
                        except:
                            pdf_text = fromatString(pdf_text)
                            index = extract_json_objects(toc_generator(pdf_text,gender))
                        for id in index:
                            indexDict.append(id)
                        pdf_text = ''
                        count = 0
            if pdf_text:
                try:
                    index = extract_json_objects(toc_generator(pdf_text,gender))
                except:
                    pdf_text = fromatString(pdf_text)
                    index = extract_json_objects(toc_generator(pdf_text,gender))
                for id in index:
                    indexDict.append(id)

        return indexDict


def extract_json_objects(text):
    pattern = r'\{[\s\S]*?\}'
    matches = re.findall(pattern, text)

    json_objects = []
    for match in matches:
        # Remove newline characters from match
        match = match.replace('\n', ' ')

        try:
            json_objects.append(json.loads(match))
        
        except SyntaxError as s:
            # st.warning(f"Error: {s}")
            try:
                stri = fromatString(f"Error: {s}\n String:{match}")
                match = re.findall(pattern, stri)
                match = match.replace('\n', ' ')
                json_objects.append(json.loads(match))
                # json_objects.append(extract_json_objects(stri)[0])
            except:
                continue
        except:
            continue
        # except json.JSONDecodeError as e:
        #     print(f"Error parsing json object: {e}")
    
    return json_objects



# *********************************************************************************

def extract_details(text):
    pattern = r'{\s*(.*?)\s*}'
    text = text.replace('\n', ' ')
    match = re.search(pattern, text, re.DOTALL)

    if match:
        dict_str = match.group(1)
        extracted_dict = {}
        try:
            extracted_dict = eval(f"""{'{'}{dict_str}{'}'} """)
        except SyntaxError as s:
            stri = fromatString(f"Error: {s}\n String:{match}")
            stri = stri.replace('\n',' ')
            match = re.search(pattern, stri, re.DOTALL)
            if match:
                dict_str = match.group(1)
                extracted_dict = eval(f"""{'{'}{dict_str}{'}'} """)
        except:
            extracted_dict = {}
        return extracted_dict
    else:
        return {}


def request_details(text):
    prompt = {'messages': [{"role": "system", "content": "You are a helper function, you will be provided with the text from deposition document. Your task is to determine witness's gender and return it in dict format."},
                         {"role": "user", "content": f"""
You are provided the deposition text delimited by ```.
Your task is to perform the following actions:
1. Determine gender of the witness male or female. Properly recognize the significance of the pronoun that will help in determining the gender of witness.
2. If you got the details format it in a python dictionary with following keys:
    "witness_gender" : ,
    "status" : True or False 
3. Return the output response in python dictionary.

Deposition text: ```{text}```
"""}]}
    
    details = get_completions(prompt)
    details_dict = extract_details(details)
    return details_dict



def most_common_word_length(word_lists):
    lengths = []
    for word_list in word_lists:
        for word in word_list:
            lengths.append(len(word))
    
    word_lengths = Counter(lengths)
    
    if not word_lengths:
        return 4
    try:
        most_common_length = word_lengths.most_common(1)[0][0]
    except:
        most_common_length = 4
    
    return most_common_length


def get_details(pdf_file):
    st.spinner("Extracting Details..")
    if pdf_file:
        uploaded_file_name = f"{pdf_file.name}_{pdf_file.size}"

        with open(f"{uploaded_file_name}.pdf", "wb") as f:
            f.write(pdf_file.read())
        
        extract_pdf_text_to_dataframe(uploaded_file_name)


        Low = []
        with open(f"{pdf_file.name}_{pdf_file.size}.pdf", "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            pdf_text = ''
            for page_num in range(0,5):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if not text:
                    # st.warning("Unable to Parse Pdf. Skipping Page")
                    continue
                else:
                    pdf_text = text
                    detail_dict = request_details(pdf_text)

                    if (("status" in list(detail_dict.keys())) and (detail_dict['status'] == "True" or detail_dict['status'] == "true" or detail_dict['status'] == True)) or (("Status" in list(detail_dict.keys())) and (detail_dict['Status'] == "True" or detail_dict['Status'] == "true" or detail_dict['Status'] == True)):
                        Low.append(list(detail_dict.values())[:-1])
        commonLength = most_common_word_length(Low)
        gender = ''
        if commonLength == 6:
            gender = "Female"
        else:
            gender = "Male"
        return gender


def normalize_text(text):
    # Normalize Text
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r". ,", "", text)
    text = text.replace("..", ".")
    text = text.replace(". .", ".")
    text = text.replace("\n", "")
    text = text.strip()

    # Tokenization
    tokens = nltk.word_tokenize(text)

    # Lowercasing
    tokens = [token.lower() for token in tokens]

    # Stopword Removal
    stop_words = set(stopwords.words('english'))
    tokens = [token for token in tokens if token not in stop_words]

    # Punctuation Removal
    tokens = [token for token in tokens if re.match(r'\w', token)]

    return " ".join(tokens)

# ******************************** MAIN *******************************************
def app1():

    st.title('DepoWizTOC_V3')


    depoDocu = st.file_uploader("Upload Deposition Document (Pdf file):", type="pdf")

    if depoDocu:
        st.success('File uploaded successfully!', icon="✅")


    if st.button('Generate TOC') and depoDocu :
        docName = str(depoDocu.name)
        gender = get_details(depoDocu)
        data = pdfParser(depoDocu,gender)

        # Create an empty table with the headers
        table_data = [['Page_no','Title','Topics','Summary']]

        for d in data:
            table_data.append(list(d.values())[:4])

        df = pd.DataFrame(table_data[1:], columns=table_data[0])

        # convert any non-string values in 'Topics' and 'Summary' columns to empty strings
        df['Topics'] = df['Topics'].apply(lambda x: str(x) if x else '')
        df['Summary'] = df['Summary'].apply(lambda x: ' '.join(x) if isinstance(x, list) else str(x))

        st.table(df)

        st.success("TOC has been Generated Successfully")
        
        # export the dataframe to a CSV file
        df_name = make_valid_partition_name(f'{depoDocu.name}_{depoDocu.size}')
        df.to_csv(f"{df_name}.csv", index=False)
        
        # concatenate the 'Title', 'Topics', and 'Summary' columns
        concatenated = df['Title'].str.cat([df['Topics'], df['Summary']], sep=' | ')

        # create a new column by merging the concatenated strings
        df['content'] = concatenated

        # drop the original 'Title', 'Topics', and 'Summary' columns
        df = df.drop(['Title', 'Topics', 'Summary'], axis=1)

        df['content']= df["content"].apply(lambda x : normalize_text(x))

        save_to_database(df_name)
        # extract_pdf_text_to_dataframe(depoDocu)

        st.table(df)


        df = pd.DataFrame(table_data[1:], columns=table_data[0])
    
        # Create a new Word document
        document = Document()

        # Add heading to the document
        heading = document.add_heading(f"{docName}", level=1)
        heading2 = document.add_heading("Table Of Content", level=2)

        # Add a table to the document and fill it with your dataframe
        table = document.add_table(rows=len(df)+1, cols=len(df.columns))
        for i in range(len(df.columns)):
            table.cell(0, i).text = df.columns[i]
        for i in range(len(df)):
            for j in range(len(df.columns)):
                table.cell(i+1, j).text = str(df.values[i, j])

        # Save the document as a bytes object
        doc_bytes = io.BytesIO()
        document.save(doc_bytes)
        doc_bytes.seek(0)
        # Download the document as a file
        b64 = base64.b64encode(doc_bytes.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="TOC_{docName}.docx">Download Deposition TOC</a>'
        st.markdown(href, unsafe_allow_html=True)
        
