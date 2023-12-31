import streamlit as st
import csv
import json
import openai
import time
import requests
import PyPDF2
import pandas as pd
import numpy as np
import re
from builtins import PendingDeprecationWarning
# from openai.embeddings_utils import get_embedding

from docx.shared import Inches
import io
from io import BytesIO
import base64
from docx import Document

from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct
from qdrant_client import models
import nltk
from nltk.corpus import stopwords





 
nltk.download('stopwords')
nltk.download('punkt')

# Global Parameters *********************************************

FILE = ''
DIMENSION = 1536  # Embeddings size
# COLLECTION_NAME = 'Deposition_Docs'  # Collection name
QDRANT_HOST = 'HOST'  # Qdrant server URI
QDRANT_PORT = 'PORT'
OPENAI_ENGINE = 'text-embedding-ada-002'  # Which engine to use

openai.api_base = "https://chat-gpt-a1.openai.azure.com/"
openai.api_type = "azure" # manually added by Vikash Sahni
openai.api_key = 'API_KEY'  # Use your own Open AI API Key here
openai.api_version = "2022-12-01" # manually added by Vikash Sahni



# connecting to the server

client = QdrantClient(host="HOST", port="PORT")


# My API Key
api_key = "API_KEY"


token = "50.0"
temp = "0.0"
top_p = "0.90"
f_pen = "0.0"
p_pen = "0.0"


# ***************************************************************


# Functions *****************************************************
def csv_load(file):
    with open(file, newline='',encoding='utf-8') as f:
        reader=csv.reader(f, delimiter=',')
        for row in reader:
            yield row



def embed(text):
    embedding = openai.Embedding.create(
        input=text,
        engine="text-embedding-ada-002"
    )
    embeddings = embedding['data'][0]['embedding']
    return embeddings


# def embed(text):
#     embedding = openai.Embed.from_prompt(text, engine=OPENAI_ENGINE)
#     return embedding["choices"][0]["embedding"]



def make_valid_partition_name(name):
  # Strip any spaces from the name.
  name = name.strip()

  # Replace any special characters with underscores.
  for char in " .,(){}[]<>*-":
    name = name.replace(char, "_")

  # Check if the name is too long.
  if len(name) > 64:
    st.warning("The file name is too long.")
    exit()

  # Return the valid partition name.
  return name





def extract_pdf_text_to_dataframe(file):
    pdf = open(f"{file}.pdf", 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf)

    extracted_text = []
    page_numbers = []

    for page_num, page in enumerate(pdf_reader.pages):
        text = page.extract_text()
        extracted_text.append(text)
        page_numbers.append(page_num+1)

    data = {'page_no': page_numbers, 'content': extracted_text}
    df = pd.DataFrame(data)
    df_name = make_valid_partition_name(file)

    df.to_csv(f"dataframe_{df_name}.csv", index=False)
    st.success("pdf to dataframe conversion successful")
    # return df_name




def normalize_text(text):
    # Normalize Text
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r". ,", "", text)
    text = text.replace("..", ".")
    text = text.replace(". .", ".")
    text = text.replace("\n", "")
    text = text.strip()

    # Tokenization
    tokens = nltk.word_tokenize(text)

    # Lowercasing
    tokens = [token.lower() for token in tokens]

    # Stopword Removal
    stop_words = set(stopwords.words('english'))
    tokens = [token for token in tokens if token not in stop_words]

    # Punctuation Removal
    tokens = [token for token in tokens if re.match(r'\w', token)]

    return " ".join(tokens)




# def append_name_to_file(file_name, name):
#     with open(file_name, 'a') as file:
#         file.write(name + '\n')



def save_to_database(dfName):
    collection = make_valid_partition_name(dfName)
    # append_name_to_file("collection.txt", collection)

    client.recreate_collection(
    collection_name=collection,
    vectors_config=models.VectorParams(size=DIMENSION, distance=models.Distance.COSINE),
    )

    FILE = dfName+".csv"
    
    points =[]

    for idx, row in enumerate(csv_load(FILE)):
        points.append(PointStruct(id=idx, vector=embed(row[1]), payload={"page_no":row[0]}))
    
    operation_info = client.upsert(
        collection_name= collection,
        wait=True,
        points=points
    )

    points = []

    print("Operation info:", operation_info)
    st.success("Dataframe saved to database. Successful")




# def fetch_documents():
#     with open("collection.txt", 'r') as file:
#         names = file.read().splitlines()
#     return names

def fetch_documents():
    collections_response = [name.name for name in client.get_collections().collections]
    return collections_response




def search(query,collection):
    # st.write(collection)
    search_result = client.search(
        collection_name=collection,
        query_vector= embed(normalize_text(query)),
        limit=20,
    )
    # st.write(search_result)
    page_list = []
    for result in search_result:
        # print(result)
        page_list.append(result.payload['page_no'])
    # st.write(page_list)

    return page_list



def get_completions(prompt, temp='0.0',top_p='0.75'):
    url = "AzureOpenAI_model_API"
    headers = {'Content-Type': 'application/json',
               'api-key': api_key,
               'max_tokens': str(token),
               'temperature': str(temp) if temp else "0.0",
               'top_p': str(top_p) if top_p else "0.75",
               'frequency_penalty': str(f_pen) if f_pen else "0.0",
               'Presence_penalty': str(p_pen) if p_pen else "0.0"
               }
    data = prompt

    body = str.encode(json.dumps(data))

    r = requests.post(url=url, data=body, headers=headers)
    response = r.content

    response = json.loads(response.decode())
    if not ('choices' in response):
        return "Choices Error"
    # print(response)
    response = response["choices"][0]["message"]["content"]
    
    return response


def getAns(pdf_text, que, name):
    data = {'messages': [{"role": "system", "content": f"""You are a deposition QnA bot,and based on the given deposition conversation text, you provide the best answer to the question asked by an attorney. Remember everywhere in the answer address the witness by their name, witness name: {name}. """},
{"role": "user", "content": f"""
You are a Deposition Document QnA bot.
Your task is to generate and give answer to the given question in a consistent style.
You are provided with the Deposition document text and question asked by an attorney.
Review the deposition text and based on you creativity and understanding, If the answer to the question can be generated from depositon text,
then generate the best possible answer to the given question. Below each answer mention the page number and line number of the discussion enclosed in parenthesis ().
If the answer cannot be generated from the deposition text,
then simply write \"Answer cannot be generated. *** \"
Generate answer in a consistent style.

<text>: ``` *** ROUGH DRAFT.  NOT CERTIFIED. ***
  1 R OUGH DRAFT DISCLAIMER
  2
  3 California Code  of Civil Procedur e Section 2025.540 
  4 states:  "When prepared as a r ough draft t ranscript, the 
  5 transcript of the deposition may not be certified and 
  6 may not be used, cite d, or transcribed as the certified 
  7 transcript of the deposition proceedings.  The rough 
  8 draft transcript m ay not be cited  or used in any way or 
  9 at any time to rebut or contradic t the certified 
 10 transcript of depo sition proceedi ngs as provi ded by the 
 11 depositio n officer."
 12
 13 This is an uncertified r ough draft that  has been 
 14 prepared in rough edi t form at counsel' s request and for 
 15 counsel's conve nience.  No represent ation is made about 
 16 its accuracy.
 17 18
 19
 20 VIDEOGRAPHER:  We are now on t he record.  This 
 21 is the videotaped deposition of Pene lope Bake r, volume 
 22 three.  Toda y's date is April 14t h, 2023.  Time is 
 23 9:13 a.m. Pacific.  A ll attorneys pr esent will be 
 24 indicated on th e stenographic record.  Please  remember 
 25 to unmute if you need to speak du ring the dep osition.  
*** ROUGH DRAFT.  NOT CERTIFIED. ***1
*** ROUGH DRAFT.  NOT CERTIFIED. ***
  1 The court will now remind t he witness they  are still 
  2 under oath.  3 PENELOPE BAKER, 
  4 the Plaintiff, called for continu ed examination by the 
  5 Defendants, bei ng first reminded to tel l the truth, the 
  6 whole truth, an d nothing but the truth, testified as 
  7 follows:
  8 MR. LOOMIS:  This is John Loom is.  It's my 
  9 understanding that we have sti pulation for a court 
 10 reporter to a effe ct that the obj ection of on e defendant 
 11 or even of a plaintiff I gu ess is of al l subject to 
 12 objecting out l ater and the same as true f or motions to 
 13 strike and r eservations of rights.  Is that correct, 
 14 sir. 15 MR. FULLE R:    Great.
 16 MR. LOOMIS:  Okay.  That's Bre tt.  All right.
 17 EXAMINATION 18 BY MR. LOOMIS:
 19 Q. All right.  Dr. Ba ker, how are you feeling 
 20 today? 21 A. Generally  not well.
 22 Q. How are y ou feeling comp ared to about a week 
 23 ago when we last talked?
 24 A. Not as well.  I 'm having more shortness of 
 25 breath and more pleurisy.
*** ROUGH DRAFT.  NOT CERTIFIED. ***2
 ```

<Question>: give medical history of witness.
 
<Answer>: Penelope Baker had a medical visit in April of 2023, which was for a cancer treatment for her mesothelioma with her oncologist, Dr. Khattab. The treatment that she received was an IV infusion with Pemetrexed, Dexamethasone, and Aloxi. She is due for another infusion on May the 9th, which will also be the same kind of infusion. Before each treatment, blood work is done at the same office and takes around 30 minutes resulting in the complete blood count. The results for the complete blood count come back in 30 minutes, and the chemistry follows that. In terms of testing, the witness is undergoing a chemistry profile, and the results usually get results within a couple of days, which is done at Sonora Quest Laboratory(Page 1, line 4-25, Page 2, line 1-25).
 
<text>: ``` *** ROUGH DRAFT.  NOT CERTIFIED. ***
  1 A. At that time he  was methodi cally disabled.
  2 Q. When did he bec ome medicall y disabled?
  3 A. To the best of my recollection  196 -- 
  4 1969-1970.
  5 Q. And how did he bec ome medically disabled, if 
  6 you know?
  7 A. He had a heart attack and d amage from the heart 
  8 attack.    9 Q. Okay.  So it's my understanding that he died in 
 10 age 70 in 1994.  That would suggest that he was born in 
 11 about 1924.  Does tha t sound a bout right?
 12 A. Correct.
 13 Q. Therefore he was a round 31 when y ou were born 
 14 and so if my math is correct, and  it often is not, that 
 15 would mean that he  had this heart at tack at a time when 
 16 you were about 14 or 15; is th at correct?
 17 A. Correct.  Pretty close.
 18 Q. So he would hav e then been approx imately 45 or 
 19 so when he had his he art attack; Is that correct?
 20 A. I'm not sure  about the age, but it was in that 
 21 time time  frame, yes.
 22 Q. Okay.  Had he been  experiencing to your 
 23 knowledge or un derstanding any problems re lated to his 
 24 heart prior to tha t heart attack?
 25 A. As a child, a very young child, he had 
*** ROUGH DRAFT.  NOT CERTIFIED. ***25
 ```

<Question>: What did Ms.Yu said about PEAKS?

<Answer>: Answer cannot be generated. ***

<text>: ``` {pdf_text} 
```
<Question>: {que}

<Answer>: 

"""}]}

    timer = 10
    for _ in range(3):
        try:
            answers = get_completions(data)
            break
        except Exception as e:
            st.write(f"Retrying after {timer} seconds...")
            time.sleep(timer)
            timer += 10
    return answers




# def getAns(pdf_text, que, name):
#     data = {'messages': [{"role": "system", "content": f"""You are a deposition QnA bot,and based on the given deposition conversation text, you provide the best answer to the question asked by an attorney. Remember everywhere in the answer address the witness by their name, witness name: {name}. """},
# {"role": "user", "content": f"""
# You are a Deposition Document QnA bot.
# You will be provided with the Deposition document text and question asked by an attorney delimited by triple backtics ```.
# If the answer to the question can be generated from depositon text,
# then review the deposition text and based on you creativity and understanding generate the best possible answer to the given question. Below each answer mention the page number and line number of the discussion enclosed in parenthesis ().
# If the answer cannot be generated from the deposition text,
# then simply write \"Answer cannot be generated. *** \"

# Provided Details:
# ```
# Question by an Attorney: {que}
# Deposition document text: {pdf_text}```
# """}]}

#     timer = 10
#     for _ in range(3):
#         try:
#             answers = get_completions(data)
#             break
#         except Exception as e:
#             st.write(f"Retrying after {timer} seconds...")
#             time.sleep(timer)
#             timer += 10
#     return answers

# def getAns(pdf_text, que, name):
#     data = {'messages': [{"role": "system", "content": f"""You are a deposition QnA bot,and based on the given deposition conversation text, you provide the best answer to the question asked by an attorney. witness name: {name}. Remember everywhere in the answer address the witness by their name. """},
# {"role": "user", "content": f"""
# You are a Deposition Document QnA bot.
# You are provided with the question asked by the attorney and the deposition document conversation text delimited by ``` .
# Your task is to review the deposition text and based on you creativity and understanding you have to provide the best possible answer to the given question.
# Below each answer mention the page number and line number of the discussion enclosed in parenthesis ().

# Important NOTE: If answer cannot be found then just return *** .

# Provided Details:
# ```
# Question by an Attorney: {que}
# Deposition document text: {pdf_text}```
# """}]}
    
# # If the answer cannot be found in Depostion document text then just return "NA".
# # (NOTE: If you couldnt find the answer in given text then just return " NA ".)

#     timer = 10
#     for _ in range(3):
#         try:
#             answers = get_completions(data)
#             break
#         except Exception as e:
#             st.write(f"Retrying after {timer} seconds...")
#             time.sleep(timer)
#             timer += 10
#     return answers



def summarizeAns(loa,query):
    data = {'messages': [{"role": "system", "content": "You are a deposition QnA bot,and based on the given list of answers and a query you create the detailed answer to the query asked by an attorney."},
{"role": "user", "content": f"""
You are a Deposition Document QnA bot.
You are provided with the query asked by the attorney and the list of answers generated by GPT delimited by ``` .
Your task is to review the answers and query and based on you creativity and understanding you have to create the best possible detailed answer to the given query.
Provided Details:
```
Query by an Attorney: {query}
List of answers: {loa}```
"""}]}

    ans = get_completions(data)
    return ans




def searchDf(listOfPgno, qnaDf, question, name):
    # st.write(listOfPgno)
    content_list = []
    answer = []
    
    for pageno in listOfPgno:
        index = pageno - 1  # Convert pageno to dataframe index
        content_list.extend(qnaDf.loc[index:index+4, 'content'].tolist())
        ans = getAns(content_list, question, name)
        content_list = []

        # st.write(len(ans))
        if len(ans)<200 or "***" in ans:
            continue
        # st.write(len(ans))

        st.write(f"Page no.{pageno}:")
        st.write(ans)
        answer.append(f"""Page no.{pageno}:\n\n {ans} \n\n ************************************** \n""")
        st.write("**************************************")
    
    return answer

def format_list(lst):
    result = []
    i = 0
    while i < len(lst):
        result.append(lst[i])
        if i+2 < len(lst) and lst[i]+2 == lst[i+2]:
            i += 2
        else:
            i += 1
    return result



def request_details(text):
    prompt = {'messages': [{"role": "system", "content": "You are a helper function, you will be provided with the text from deposition document. Your task is to extract and return the witness's name."},
                         {"role": "user", "content": f"""
You are provided the deposition text delimited by ```.
Your task is to perform the following actions:
1. Extract the Full name of the witness.
Deposition text: ```{text}```
"""}]}

    details = get_completions(prompt)
    return details

def extract_name_from_csv(df):
    content_list = df['content'].head(5).tolist()
    name = request_details(content_list)
    return name

# ***************************************************************




def app2():

    # Sidebar
    # drop = st.sidebar.button("Drop Collection")
    # if drop:
    #     collectionsList = fetch_documents()
    #     for collection in collectionsList:
    #         client.delete_collection(collection_name=collection)
    #         with open('collection.txt', 'w') as file:
    #             file.truncate(0)
    #     st.success(f"The file collections has been cleared.")


    st.sidebar.markdown("---")
    st.sidebar.subheader("Select Document for QnA")
    document_list = st.sidebar.selectbox("Documents", fetch_documents())
    # save_preferences = st.sidebar.button("Save Preference")
    # st.write(document_list)




    # Main screen
    st.title("Talk with your document")
    # st.subheader("Talk to Document: Question and Answering over Document")

    with st.expander("Follow these steps to use the app"):
        # st.markdown(":")
        # st.markdown("1. Upload the document in Sidebar uploader")
        st.markdown("1. Select the documents for QnA in sidebar")
        st.markdown("2. Search for a query and get response")
        # Add more tutorial steps as needed

    que = st.text_input("Enter your query or question")
    get_response_button = st.button("Get Response")


    if get_response_button and que:
        # utility.drop_collection(COLLECTION_NAME)
        pglist = search(que,document_list)
        loi = []
        for integer in pglist:
            if integer.isdigit():
                loi.append(int(integer))
        loi.sort()
        loi = format_list(list(set(loi)))
        # st.write(loi)
        # print(loi)
        # import the CSV file as a dataframe
        st.write(f"Selected Document: {document_list}")
        qnaDf = pd.read_csv(f"dataframe_{document_list}.csv")

        witness_name = extract_name_from_csv(qnaDf)
        # st.write(witness_name)
        
        
        answer = searchDf(loi, qnaDf, que, witness_name)


        st.write(que)
        ans = summarizeAns("\n".join(answer),que)
        st.write("### Summarized Answer")
        st.write(ans)

        document = Document()

        answers = "\n\n".join(answer)

        # Add heading to the document
        heading = document.add_heading("Generated Response", level=1)
        document.add_paragraph(f"Question: {que}")
        heading2 = document.add_heading("Answers", level=2)
        document.add_paragraph(answers)
        heading2 = document.add_heading("Summarized Answer", level=2)
        document.add_paragraph(ans)
        

        # Save the document as a bytes object
        doc_bytes = io.BytesIO()
        document.save(doc_bytes)
        doc_bytes.seek(0)
        # Download the document as a file
        b64 = base64.b64encode(doc_bytes.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="QNA_{document_list}.docx">Download QNA Docx</a>'
        st.markdown(href, unsafe_allow_html=True)
        
        # Handle getting response for the query
